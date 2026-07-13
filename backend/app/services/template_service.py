import os
import uuid
import shutil
from datetime import datetime
from typing import List, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.config import settings
from app.models.schemas import TemplateCreate, TemplateResponse, DocumentType


class TemplateService:
    def __init__(self):
        self.template_dir = settings.TEMPLATE_DIR
        os.makedirs(self.template_dir, exist_ok=True)
        # 模拟数据库存储
        self._templates: dict = {}
    
    def save_template(self, file_path: str, template_data: TemplateCreate) -> TemplateResponse:
        """保存上传的模板文件"""
        template_id = str(uuid.uuid4())
        ext = os.path.splitext(file_path)[1]
        new_filename = f"{template_id}{ext}"
        dest_path = os.path.join(self.template_dir, new_filename)
        
        # 复制文件到模板目录
        shutil.copy(file_path, dest_path)
        
        template = TemplateResponse(
            id=template_id,
            name=template_data.name,
            description=template_data.description,
            document_type=template_data.document_type,
            file_path=dest_path,
            created_at=datetime.now(),
            updated_at=None
        )
        
        self._templates[template_id] = template
        return template
    
    def get_template(self, template_id: str) -> Optional[TemplateResponse]:
        """获取模板信息"""
        return self._templates.get(template_id)
    
    def list_templates(self, document_type: Optional[DocumentType] = None) -> List[TemplateResponse]:
        """列出所有模板"""
        templates = list(self._templates.values())
        if document_type:
            templates = [t for t in templates if t.document_type == document_type]
        return templates
    
    def delete_template(self, template_id: str) -> bool:
        """删除模板"""
        template = self._templates.get(template_id)
        if not template:
            return False
        
        # 删除文件
        if os.path.exists(template.file_path):
            os.remove(template.file_path)
        
        del self._templates[template_id]
        return True
    
    def extract_placeholders(self, template_id: str) -> List[str]:
        """从模板中提取占位符（如 {{field_name}}）"""
        template = self._templates.get(template_id)
        if not template or not os.path.exists(template.file_path):
            return []
        
        doc = Document(template.file_path)
        placeholders = set()
        
        # 遍历所有段落
        for para in doc.paragraphs:
            self._extract_from_text(para.text, placeholders)
        
        # 遍历所有表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._extract_from_text(para.text, placeholders)
        
        return sorted(list(placeholders))
    
    def _extract_from_text(self, text: str, placeholders: set):
        """从文本中提取 {{placeholder}} 格式的占位符"""
        import re
        pattern = r'\{\{([^}]+)\}\}'
        matches = re.findall(pattern, text)
        placeholders.update(matches)
    
    def fill_template(self, template_id: str, data: dict, output_path: str) -> str:
        """填充模板并生成文档"""
        template = self._templates.get(template_id)
        if not template or not os.path.exists(template.file_path):
            raise ValueError("模板不存在")
        
        # 加载模板文档
        doc = Document(template.file_path)
        
        # 替换段落中的占位符
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, data)
        
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, data)
        
        # 保存生成的文档
        doc.save(output_path)
        return output_path
    
    def _replace_in_paragraph(self, para, data: dict):
        """在段落中替换占位符"""
        for run in para.runs:
            text = run.text
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in text:
                    text = text.replace(placeholder, str(value))
            run.text = text


template_service = TemplateService()
