"""
历史沿革Word文档生成器
支持模板占位符替换和默认生成两种模式

模板占位符说明：
- {{company_name}}: 公司名称
- {{history_content}}: 历史沿革正文（自动解析为段落和表格）
- {{report_date}}: 报告生成日期
- {{change_count}}: 变更次数
- {{law_firm_name}}: 律所名称（可选）
- {{lawyer_name}}: 律师姓名（可选）

使用示例：
1. 无模板（默认生成）: generate_history_word_document(company_name, history_result, output_path)
2. 使用模板: generate_from_template(company_name, history_result, template_path, output_path)
"""
import os
import re
from datetime import datetime
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class HistoryDocxGenerator:
    """历史沿革Word文档生成器"""
    
    def __init__(self):
        self.default_font = '宋号'
        self.default_font_size = Pt(12)
    
    def generate_from_template(
        self,
        company_name: str,
        history_result: Dict[str, Any],
        template_path: str,
        output_path: str,
        extra_data: Dict[str, str] = None
    ) -> str:
        """
        使用模板生成Word文档（占位符替换模式）
        
        Args:
            company_name: 公司名称
            history_result: 历史沿革生成结果
            template_path: 模板文件路径(.docx)
            output_path: 输出文件路径
            extra_data: 额外数据（律所名、律师名等）
        
        Returns:
            输出文件路径
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")
        
        # 加载模板
        doc = Document(template_path)
        
        # 准备替换数据
        history_text = history_result.get('text', '')
        changes = history_result.get('changes', [])
        
        data = {
            'company_name': company_name,
            'report_date': datetime.now().strftime('%Y年%m月%d日'),
            'change_count': str(len(changes)),
            'history_content': history_text,  # 特殊处理：需要解析为段落
        }
        
        # 合并额外数据
        if extra_data:
            data.update(extra_data)
        
        # 替换所有段落中的占位符
        for para in doc.paragraphs:
            self._replace_placeholders_in_paragraph(para, data, doc)
        
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_placeholders_in_paragraph(para, data, doc)
        
        # 保存文档
        doc.save(output_path)
        return output_path
    
    def _replace_placeholders_in_paragraph(
        self, 
        para, 
        data: Dict[str, str],
        doc: Document
    ):
        """替换段落中的占位符"""
        full_text = para.text
        
        # 检查是否有特殊占位符 {{history_content}}
        if '{{history_content}}' in full_text:
            # 清空原段落
            para.clear()
            # 添加历史沿革内容（解析为段落和表格）
            self._add_history_content_to_paragraph(para, data.get('history_content', ''), doc)
            return
        
        # 普通占位符替换
        has_placeholder = False
        for key, value in data.items():
            placeholder = f'{{{{{key}}}}}'
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value))
                has_placeholder = True
        
        # 如果有替换，更新段落
        if has_placeholder and full_text != para.text:
            para.clear()
            run = para.add_run(full_text)
            self._set_run_font(run, font_size=Pt(12))
    
    def _add_history_content_to_paragraph(self, para, history_text: str, doc: Document):
        """将历史沿革内容解析为段落和表格"""
        # 由于python-docx限制，我们需要在段落位置插入多个元素
        # 这里采用简单策略：保留段落，添加文本
        
        lines = history_text.split('\n')
        current_text = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 如果是表格行，特殊标记
            if line.startswith('|'):
                # 先输出累积的文本
                if current_text:
                    para.add_run('\n'.join(current_text))
                    current_text = []
                # 表格需要后续处理，这里简化处理
                continue
            
            # 处理标题行
            if line.startswith('【') and line.endswith('】'):
                if current_text:
                    para.add_run('\n'.join(current_text) + '\n')
                    current_text = []
                run = para.add_run(line + '\n')
                self._set_run_font(run, font_size=Pt(14), bold=True)
            else:
                current_text.append(line)
        
        # 输出剩余文本
        if current_text:
            para.add_run('\n'.join(current_text))
    
    def generate_history_document(
        self,
        company_name: str,
        history_text: str,
        changes: List[Dict],
        output_path: str
    ):
        """
        生成历史沿革Word文档（默认模式，无模板）
        
        Args:
            company_name: 公司名称
            history_text: 历史沿革文本（Markdown格式）
            changes: 变更记录列表
            output_path: 输出文件路径
        """
        doc = Document()
        
        # 设置文档默认字体
        self._set_document_default_font(doc)
        
        # 添加标题
        title = doc.add_heading(f"{company_name}历史沿革", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_run_font(title.runs[0], font_size=Pt(18), bold=True)
        
        # 解析并添加历史沿革内容
        self._add_history_content(doc, history_text)
        
        # 保存文档
        doc.save(output_path)
    
    def _set_document_default_font(self, doc: Document):
        """设置文档默认字体"""
        style = doc.styles['Normal']
        font = style.font
        font.name = self.default_font
        font.size = self.default_font_size
        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
    
    def _set_run_font(self, run, font_name: str = None, font_size: Pt = None, bold: bool = False):
        """设置文本run的字体"""
        if font_name:
            run.font.name = font_name
            run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if font_size:
            run.font.size = font_size
        run.font.bold = bold
    
    def _add_history_content(self, doc: Document, history_text: str):
        """添加历史沿革内容（解析Markdown格式）"""
        lines = history_text.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # 跳过空行
            if not line:
                i += 1
                continue
            
            # 处理序号标题行：【2025年10月11日 股权转让】
            if line.startswith('【') and line.endswith('】'):
                heading_text = line[1:-1]  # 去掉【】
                heading = doc.add_heading(heading_text, level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if heading.runs:
                    self._set_run_font(heading.runs[0], font_size=Pt(14), bold=True)
                i += 1
                continue
            
            # 处理表格开始
            if line.startswith('| 序号 |') or line.startswith('|:--:|'):
                # 收集表格所有行
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                
                if table_lines:
                    self._add_shareholder_table(doc, table_lines)
                continue
            
            # 处理普通段落
            if line:
                para = doc.add_paragraph(line)
                para.paragraph_format.line_spacing = 1.5
                para.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
                for run in para.runs:
                    self._set_run_font(run, font_size=Pt(12))
            
            i += 1
    
    def _add_shareholder_table(self, doc: Document, table_lines: List[str]):
        """添加股权结构表格"""
        if len(table_lines) < 2:
            return
        
        # 解析表头
        headers = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
        
        # 解析数据行（跳过表头和对齐行）
        data_rows = []
        for line in table_lines[2:]:
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                data_rows.append(cells)
        
        if not data_rows:
            return
        
        # 添加表格说明文字
        intro_para = doc.add_paragraph("本次增资和股权转让完成后，公司的股权结构如下：")
        intro_para.paragraph_format.line_spacing = 1.5
        for run in intro_para.runs:
            self._set_run_font(run, font_size=Pt(12))
        
        # 创建表格
        table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表格宽度
        table.autofit = False
        table.allow_autofit = False
        
        # 添加表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    self._set_run_font(run, font_size=Pt(12), bold=True)
            # 设置背景色
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D9D9D9')
            cell._element.get_or_add_tcPr().append(shading)
        
        # 添加数据行
        for row_idx, row_data in enumerate(data_rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(row_cells):
                    cell = row_cells[col_idx]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            self._set_run_font(run, font_size=Pt(12))
        
        # 设置列宽
        self._set_table_column_widths(table)
        
        # 添加表格后空行
        doc.add_paragraph()
    
    def _set_table_column_widths(self, table):
        """设置表格列宽"""
        widths = [1, 4, 2, 2]
        total_width = sum(widths)
        
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx < len(widths):
                    cell.width = Cm(15 * widths[idx] / total_width)


# 模板存储目录
TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), '..', '..', 'templates', 'word')
os.makedirs(TEMPLATE_DIR, exist_ok=True)


def get_template_path(template_name: str = None) -> Optional[str]:
    """获取模板文件路径"""
    if template_name:
        path = os.path.join(TEMPLATE_DIR, template_name)
        if os.path.exists(path):
            return path
    
    # 查找默认模板
    default_templates = ['历史沿革模板.docx', 'template.docx', 'default.docx']
    for tmpl in default_templates:
        path = os.path.join(TEMPLATE_DIR, tmpl)
        if os.path.exists(path):
            return path
    
    return None


def list_available_templates() -> List[str]:
    """列出所有可用模板"""
    if not os.path.exists(TEMPLATE_DIR):
        return []
    
    templates = []
    for f in os.listdir(TEMPLATE_DIR):
        if f.endswith('.docx'):
            templates.append(f)
    return templates


def generate_history_word_document(
    company_name: str,
    history_result: Dict[str, Any],
    output_path: str,
    template_path: str = None,
    extra_data: Dict[str, str] = None
) -> str:
    """
    生成历史沿革Word文档的便捷函数
    
    Args:
        company_name: 公司名称
        history_result: 历史沿革生成结果
        output_path: 输出文件路径
        template_path: 模板文件路径（可选，不传则使用默认生成）
        extra_data: 额外数据（可选）
    
    Returns:
        输出文件路径
    """
    generator = HistoryDocxGenerator()
    
    # 如果指定了模板且存在，使用模板模式
    if template_path and os.path.exists(template_path):
        return generator.generate_from_template(
            company_name=company_name,
            history_result=history_result,
            template_path=template_path,
            output_path=output_path,
            extra_data=extra_data
        )
    
    # 否则使用默认生成模式
    history_text = history_result.get('text', '')
    changes = history_result.get('changes', [])
    
    generator.generate_history_document(
        company_name=company_name,
        history_text=history_text,
        changes=changes,
        output_path=output_path
    )
    
    return output_path


if __name__ == "__main__":
    # 测试
    test_result = {
        "text": """【2025年10月11日 股权转让、增资】

2025年10月11日，公司召开股东会并作出决议：

公司注册资本由人民币2740.286200万元变更为人民币3052.346700万元。

本次增资和股权转让完成后，公司的股权结构如下：

| 序号 | 股东名称/姓名 | 出资额（万元） | 出资比例 |
|:---:|:---:|:---:|:---:|
| 1 | 张新峰 | 【**】 | 【**】% |
| 合计 | - | 【**】 | 100.0000% |

2025年10月11日，公司完成了上述变更的工商变更登记程序。""",
        "changes": []
    }
    
    import tempfile
    output = tempfile.mktemp(suffix='.docx')
    generate_history_word_document("测试公司", test_result, output)
    print(f"文档已生成: {output}")
