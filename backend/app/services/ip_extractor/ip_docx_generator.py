"""
知识产权 Word 文档生成器

格式与历史沿革 Word 保持一致（继承 HistoryDocxGenerator 的字体/段落规范）：
- 标题：四号加粗居中
- 正文：小四，1.25 倍行距，段前段后各 0.5 行
- 中文宋体，英文/数字 Times New Roman

结构：标题 → 一、商标（汇总段 + 商标表，图案列插图）→ 二、专利（汇总段 + 专利表）
"""
from io import BytesIO
from typing import Dict, List

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.services.qcc_extractor.history_docx_generator import HistoryDocxGenerator
from .service import IPExtractionResult


class IPDocxGenerator(HistoryDocxGenerator):
    """知识产权 Word 文档生成器（沿用历史沿革的格式规范）"""

    def generate(self, result: IPExtractionResult, output_path: str) -> str:
        doc = Document()
        self._set_document_default_font(doc)

        # 标题（四号加粗居中）
        title = doc.add_heading(f"{result.company_name}知识产权", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_run_font(title.runs[0], font_size=self.TITLE_SIZE, bold=True)
        title.runs[0].font.color.rgb = RGBColor(0, 0, 0)

        # 一、商标
        self._add_section_heading(doc, "一、商标")
        self._add_body_paragraph(doc, result.trademark_summary_text())
        if result.trademarks:
            self._add_trademark_table(doc, result.trademarks, result.trademark_images)
        else:
            self._add_body_paragraph(doc, "【**】")

        # 二、专利
        self._add_section_heading(doc, "二、专利")
        self._add_body_paragraph(doc, result.patent_summary_text())
        if result.patents:
            self._add_patent_table(doc, result.patents)
        else:
            self._add_body_paragraph(doc, "【**】")

        doc.save(output_path)
        return output_path

    # ----------------------------------------------------------------

    def _add_section_heading(self, doc: Document, text: str):
        heading = doc.add_heading(text, level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if heading.runs:
            self._set_run_font(heading.runs[0], font_size=self.TITLE_SIZE, bold=True)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        self._set_para_format(heading)

    def _add_body_paragraph(self, doc: Document, text: str):
        para = doc.add_paragraph(text)
        self._set_para_format(para)
        for run in para.runs:
            self._set_run_font(run, font_size=self.BODY_SIZE)

    def _new_table(self, doc: Document, headers: List[str], n_rows: int):
        table = doc.add_table(rows=n_rows + 1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    self._set_run_font(run, font_size=self.BODY_SIZE, bold=True)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D9D9D9')
            cell._element.get_or_add_tcPr().append(shading)
        return table

    def _fill_cell(self, cell, text: str):
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                self._set_run_font(run, font_size=self.BODY_SIZE)

    def _set_col_widths(self, table, widths_cm: List[float]):
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx < len(widths_cm):
                    cell.width = Cm(widths_cm[idx])

    def _add_trademark_table(self, doc: Document, trademarks: List[Dict],
                             images: Dict[str, bytes]):
        headers = ["序号", "商标图案", "商标名称", "商标状态", "申请/注册号", "申请日期", "国际分类"]
        table = self._new_table(doc, headers, len(trademarks))
        for row_idx, tm in enumerate(trademarks):
            cells = table.rows[row_idx + 1].cells
            self._fill_cell(cells[0], str(tm["seq"]))
            # 图案列：有图插图，无图占位
            img = images.get(tm["app_no"])
            if img:
                para = cells[1].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    para.add_run().add_picture(BytesIO(img), width=Cm(2))
                except Exception:
                    self._fill_cell(cells[1], "【**】")
            else:
                self._fill_cell(cells[1], "【**】")
            self._fill_cell(cells[2], tm["name"])
            self._fill_cell(cells[3], tm["status"])
            self._fill_cell(cells[4], tm["app_no"])
            self._fill_cell(cells[5], tm["app_date"])
            self._fill_cell(cells[6], tm["intl_class"])
        self._set_col_widths(table, [1.2, 2.6, 2.6, 3.2, 2.4, 2.0, 1.6])
        doc.add_paragraph()

    def _add_patent_table(self, doc: Document, patents: List[Dict]):
        headers = ["序号", "名称", "专利类型", "法律状态", "申请号", "申请日期"]
        table = self._new_table(doc, headers, len(patents))
        for row_idx, p in enumerate(patents):
            cells = table.rows[row_idx + 1].cells
            self._fill_cell(cells[0], str(p["seq"]))
            # 名称较长，左对齐
            self._fill_cell(cells[1], p.get("name", ""))
            for paragraph in cells[1].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self._fill_cell(cells[2], p["patent_type"])
            self._fill_cell(cells[3], p["legal_status"])
            self._fill_cell(cells[4], p["app_no"])
            self._fill_cell(cells[5], p["app_date"])
        self._set_col_widths(table, [1.0, 4.8, 1.8, 1.8, 3.4, 2.2])
        doc.add_paragraph()


def generate_ip_word_document(result: IPExtractionResult, output_path: str) -> str:
    """模块级入口：生成知识产权 Word 文档"""
    return IPDocxGenerator().generate(result, output_path)
