import os
import uuid
from datetime import datetime
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.config import settings
from app.services.template_service import template_service
from app.services.ai_service import ai_service
from app.services.knowledge_service import knowledge_service


class DocumentService:
    def __init__(self):
        self.generated_dir = settings.GENERATED_DIR
        os.makedirs(self.generated_dir, exist_ok=True)
    
    async def generate_document(
        self,
        template_id: str,
        form_data: Dict[str, Any],
        use_knowledge_base: bool = True
    ) -> str:
        """生成文档"""
        
        # 获取模板
        template = template_service.get_template(template_id)
        if not template:
            raise ValueError("模板不存在")
        
        # 获取知识库上下文
        context = None
        if use_knowledge_base:
            # 构建查询语句
            query = f"{template.document_type} " + " ".join([str(v) for v in form_data.values()])
            context = knowledge_service.get_relevant_context(
                query=query,
                document_type=template.document_type
            )
        
        # 使用AI生成内容
        ai_content = await ai_service.generate_document_content(
            document_type=template.document_type,
            form_data=form_data,
            context=context
        )
        
        # 生成文档ID和文件路径
        document_id = str(uuid.uuid4())
        output_path = os.path.join(self.generated_dir, f"{document_id}.docx")
        
        # 如果模板存在且有占位符，使用模板填充
        if template and os.path.exists(template.file_path):
            try:
                self._generate_with_template(
                    template_path=template.file_path,
                    form_data=form_data,
                    ai_content=ai_content,
                    output_path=output_path
                )
            except Exception as e:
                # 如果模板处理失败，直接生成新文档
                self._generate_new_document(
                    title=template.name,
                    content=ai_content,
                    output_path=output_path
                )
        else:
            # 没有模板，直接生成新文档
            self._generate_new_document(
                title=template.name,
                content=ai_content,
                output_path=output_path
            )
        
        return document_id, output_path
    
    def _generate_with_template(
        self,
        template_path: str,
        form_data: Dict[str, Any],
        ai_content: str,
        output_path: str
    ):
        """使用模板生成文档"""
        doc = Document(template_path)
        
        # 准备填充数据
        fill_data = form_data.copy()
        fill_data['ai_generated_content'] = ai_content
        
        # 替换段落
        for para in doc.paragraphs:
            self._replace_placeholders_in_paragraph(para, fill_data)
        
        # 替换表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_placeholders_in_paragraph(para, fill_data)
        
        doc.save(output_path)
    
    def _replace_placeholders_in_paragraph(self, para, data: Dict[str, Any]):
        """替换段落中的占位符"""
        full_text = para.text
        
        # 检查是否有占位符
        if '{{' not in full_text:
            return
        
        # 替换所有占位符
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value))
        
        # 处理 {{ai_content}} 特殊占位符，将其替换为多段落
        if '{{ai_generated_content}}' in full_text:
            full_text = full_text.replace('{{ai_generated_content}}', data.get('ai_generated_content', ''))
        
        # 如果文本有变化，更新段落
        if full_text != para.text:
            # 清空原段落
            para.clear()
            # 添加新文本
            run = para.add_run(full_text)
            run.font.size = Pt(12)
    
    def _generate_new_document(self, title: str, content: str, output_path: str):
        """生成新文档（无模板）"""
        doc = Document()
        
        # 添加标题
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加内容
        # 按段落分割AI生成的内容
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # 检测标题
            if para_text.startswith('# '):
                doc.add_heading(para_text[2:], level=1)
            elif para_text.startswith('## '):
                doc.add_heading(para_text[3:], level=2)
            elif para_text.startswith('### '):
                doc.add_heading(para_text[4:], level=3)
            else:
                # 普通段落
                para = doc.add_paragraph(para_text)
                para.paragraph_format.line_spacing = 1.5
                para.paragraph_format.first_line_indent = Pt(24)
        
        doc.save(output_path)
    
    def get_document_path(self, document_id: str) -> str:
        """获取文档路径"""
        return os.path.join(self.generated_dir, f"{document_id}.docx")
    
    def delete_document(self, document_id: str) -> bool:
        """删除文档"""
        path = self.get_document_path(document_id)
        if os.path.exists(path):
            os.remove(path)
            return True
        return False


document_service = DocumentService()
