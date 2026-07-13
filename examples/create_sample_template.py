#!/usr/bin/env python3
"""
创建示例模板文件
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def create_legal_opinion_template():
    """创建法律意见书模板"""
    doc = Document()
    
    # 标题
    title = doc.add_heading("关于{{company_name}}申请首次公开发行股票并在科创板上市的", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading("法律意见书", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 致
    doc.add_paragraph("致：{{company_name}}")
    doc.add_paragraph()
    
    # 引言
    intro = doc.add_paragraph(
        "{{law_firm_name}}（以下简称\"本所\"）接受{{company_name}}（以下简称\"发行人\"或\"公司\"）的委托，"
        "担任发行人申请首次公开发行人民币普通股股票并在科创板上市（以下简称\"本次发行上市\"）的专项法律顾问。"
    )
    intro.paragraph_format.first_line_indent = Pt(24)
    
    doc.add_paragraph()
    
    # AI 生成内容占位符
    doc.add_heading("正文", level=2)
    content_para = doc.add_paragraph("{{ai_generated_content}}")
    content_para.paragraph_format.first_line_indent = Pt(24)
    
    doc.add_paragraph()
    
    # 结尾
    conclusion = doc.add_paragraph(
        "本法律意见书一式{{copies}}份，经本所负责人及经办律师签字并加盖本所公章后生效。"
    )
    conclusion.paragraph_format.first_line_indent = Pt(24)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 签署
    sign_table = doc.add_table(rows=3, cols=2)
    sign_table.cell(0, 0).text = ""
    sign_table.cell(0, 1).text = "{{law_firm_name}}（盖章）"
    sign_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    sign_table.cell(1, 0).text = ""
    sign_table.cell(1, 1).text = "负责人：{{responsible_lawyer}}"
    sign_table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    sign_table.cell(2, 0).text = ""
    sign_table.cell(2, 1).text = "经办律师：{{handling_lawyer}}"
    sign_table.cell(2, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 日期
    doc.add_paragraph()
    date_para = doc.add_paragraph("{{date}}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 保存
    output_path = "法律意见书模板.docx"
    doc.save(output_path)
    print(f"✓ 已创建: {output_path}")
    return output_path


def create_board_rules_template():
    """创建三会制度模板"""
    doc = Document()
    
    # 标题
    title = doc.add_heading("{{company_name}}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading("{{rule_name}}", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 章节
    doc.add_heading("第一章 总则", level=2)
    
    article1 = doc.add_paragraph("第一条 为规范{{company_name}}（以下简称\"公司\"）的{{rule_type}}，完善公司法人治理结构，根据《中华人民共和国公司法》等相关法律法规及《公司章程》的规定，制定本制度。")
    article1.paragraph_format.first_line_indent = Pt(24)
    
    doc.add_paragraph()
    
    doc.add_heading("正文内容", level=2)
    content_para = doc.add_paragraph("{{ai_generated_content}}")
    content_para.paragraph_format.first_line_indent = Pt(24)
    
    doc.add_paragraph()
    
    # 附则
    doc.add_heading("附则", level=2)
    final = doc.add_paragraph("本制度自发布之日起施行。")
    final.paragraph_format.first_line_indent = Pt(24)
    
    # 签署
    doc.add_paragraph()
    doc.add_paragraph()
    sign = doc.add_paragraph("{{company_name}}（盖章）")
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    date_para = doc.add_paragraph("{{date}}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 保存
    output_path = "三会制度模板.docx"
    doc.save(output_path)
    print(f"✓ 已创建: {output_path}")
    return output_path


def create_work_report_template():
    """创建律师工作报告模板"""
    doc = Document()
    
    # 标题
    title = doc.add_heading("{{law_firm_name}}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading("关于{{company_name}}申请首次公开发行股票并在科创板上市的律师工作报告", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 致
    doc.add_paragraph("致：{{company_name}}")
    doc.add_paragraph()
    
    # 引言
    intro = doc.add_paragraph(
        "{{law_firm_name}}（以下简称\"本所\"）接受{{company_name}}（以下简称\"发行人\"）的委托，"
        "担任发行人申请首次公开发行人民币普通股股票并在科创板上市（以下简称\"本次发行上市\"）的专项法律顾问。"
        "本所根据《中华人民共和国公司法》《中华人民共和国证券法》《科创板首次公开发行股票注册管理办法（试行）》"
        "等法律、法规和规范性文件的有关规定，按照律师行业公认的业务标准、道德规范和勤勉尽责精神，"
        "出具本律师工作报告。"
    )
    intro.paragraph_format.first_line_indent = Pt(24)
    
    doc.add_paragraph()
    
    # 释义
    doc.add_heading("释义", level=2)
    doc.add_paragraph("在本律师工作报告中，除非文义另有所指，下列词语具有下述涵义：").paragraph_format.first_line_indent = Pt(24)
    
    doc.add_paragraph()
    
    # AI 生成内容
    doc.add_heading("正文", level=2)
    content_para = doc.add_paragraph("{{ai_generated_content}}")
    content_para.paragraph_format.first_line_indent = Pt(24)
    
    doc.add_paragraph()
    
    # 结尾
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 签署
    sign_table = doc.add_table(rows=3, cols=2)
    sign_table.cell(0, 0).text = ""
    sign_table.cell(0, 1).text = "{{law_firm_name}}（盖章）"
    sign_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    sign_table.cell(1, 0).text = ""
    sign_table.cell(1, 1).text = "负责人：{{responsible_lawyer}}"
    sign_table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    sign_table.cell(2, 0).text = ""
    sign_table.cell(2, 1).text = "经办律师：{{handling_lawyer}}"
    sign_table.cell(2, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    date_para = doc.add_paragraph("{{date}}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 保存
    output_path = "律师工作报告模板.docx"
    doc.save(output_path)
    print(f"✓ 已创建: {output_path}")
    return output_path


if __name__ == "__main__":
    print("=" * 50)
    print("  创建示例模板文件")
    print("=" * 50)
    print()
    
    create_legal_opinion_template()
    create_board_rules_template()
    create_work_report_template()
    
    print()
    print("=" * 50)
    print("  模板创建完成！")
    print("=" * 50)
    print()
    print("使用说明：")
    print("1. 这些模板可以在 模板管理 页面上传")
    print("2. 使用 {{占位符}} 格式的内容会被自动识别为表单字段")
    print("3. {{ai_generated_content}} 将被 AI 生成的内容替换")
    print()
